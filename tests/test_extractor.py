"""Tests para la extracción de texto (PDF/Word/txt)."""
import io

from docx import Document

from app.services.ingestion.extractor import extract_text


def test_extract_txt() -> None:
    text = extract_text(b"Hola mundo", "nota.txt")
    assert text == "Hola mundo"


def test_extract_docx_paragraphs() -> None:
    doc = Document()
    doc.add_paragraph("Contrato de arrendamiento")
    doc.add_paragraph("Arrendatario: Juan Pérez")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text(buf.getvalue(), "contrato.docx")
    assert "Contrato de arrendamiento" in text
    assert "Arrendatario: Juan Pérez" in text


def test_extract_docx_tables() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Concepto"
    table.rows[0].cells[1].text = "Monto"
    table.rows[1].cells[0].text = "Renta"
    table.rows[1].cells[1].text = "$650.000"
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text(buf.getvalue(), "balance.docx")
    assert "Concepto | Monto" in text
    assert "Renta | $650.000" in text


def test_extract_xlsx() -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Balance"
    ws.append(["Cliente", "RUT", "Monto"])
    ws.append(["Juan Pérez", "12.345.678-9", 650000])
    buf = io.BytesIO()
    wb.save(buf)

    text = extract_text(buf.getvalue(), "balance.xlsx")
    assert "=== Hoja: Balance ===" in text
    assert "Cliente | RUT | Monto" in text
    assert "Juan Pérez | 12.345.678-9 | 650000" in text


def test_extract_xlsx_numeric_rut() -> None:
    """RUT guardado como número con formato (caso reportado por Codex).

    Excel auto-convierte "12.345.678-9" a 123456789 cuando la celda no es texto.
    El extractor debe reconstruir la forma con puntos y guión para que el
    scrubber de PII lo detecte correctamente.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Nómina"
    ws.cell(row=1, column=1, value="Nombre")
    ws.cell(row=1, column=2, value="RUT")
    ws.cell(row=2, column=1, value="Ana González")
    rut_cell = ws.cell(row=2, column=2, value=123456789)
    rut_cell.number_format = "##.###.###-0"  # formato RUT en Excel
    buf = io.BytesIO()
    wb.save(buf)

    text = extract_text(buf.getvalue(), "nomina.xlsx")
    # Debe verse como RUT formateado, no como número crudo
    assert "12.345.678-9" in text
    assert "123456789" not in text
